import io
import csv
from typing import Optional
from werkzeug.datastructures import FileStorage
import pdfplumber
from docx import Document
import openpyxl
from utils.validators import InputValidator

class FileProcessor:
    """File processing utilities for extracting text from various formats"""
    
    def __init__(self):
        self.validator = InputValidator()
        self.max_pages = 100  # Limit PDF pages to prevent memory issues
        self.max_rows = 10000  # Limit Excel rows
    
    def extract_text(self, file: FileStorage) -> Optional[str]:
        """Extract text from uploaded file based on file type"""
        if not self.validator.validate_file(file):
            return None
        
        filename = file.filename.lower()
        
        try:
            if filename.endswith('.pdf'):
                return self._extract_from_pdf(file)
            elif filename.endswith(('.docx', '.doc')):
                return self._extract_from_docx(file)
            elif filename.endswith(('.xlsx', '.xls')):
                return self._extract_from_excel(file)
            elif filename.endswith(('.txt', '.csv')):
                return self._extract_from_text(file)
            else:
                return None
        except Exception as e:
            print(f"Error processing file {filename}: {str(e)}")
            return None
    
    def _extract_from_pdf(self, file: FileStorage) -> Optional[str]:
        """Extract text from PDF file"""
        try:
            text_content = []
            
            with pdfplumber.open(file.stream) as pdf:
                # Limit number of pages to prevent memory issues
                pages_to_process = min(len(pdf.pages), self.max_pages)
                
                for i in range(pages_to_process):
                    page = pdf.pages[i]
                    text = page.extract_text()
                    if text:
                        text_content.append(text)
            
            return '\n'.join(text_content) if text_content else None
            
        except Exception as e:
            print(f"Error extracting PDF: {str(e)}")
            return None
    
    def _extract_from_docx(self, file: FileStorage) -> Optional[str]:
        """Extract text from DOCX file"""
        try:
            doc = Document(file.stream)
            text_content = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text)
                    if row_text:
                        text_content.append(' | '.join(row_text))
            
            return '\n'.join(text_content) if text_content else None
            
        except Exception as e:
            print(f"Error extracting DOCX: {str(e)}")
            return None
    
    def _extract_from_excel(self, file: FileStorage) -> Optional[str]:
        """Extract text from Excel file"""
        try:
            workbook = openpyxl.load_workbook(file.stream, read_only=True)
            text_content = []
            
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                
                # Add sheet name as header
                text_content.append(f"=== {sheet_name} ===")
                
                row_count = 0
                for row in sheet.iter_rows(values_only=True):
                    if row_count >= self.max_rows:
                        text_content.append("... (truncated due to size limit)")
                        break
                    
                    # Filter out None values and convert to strings
                    row_values = [str(cell) for cell in row if cell is not None and str(cell).strip()]
                    
                    if row_values:
                        text_content.append(' | '.join(row_values))
                    
                    row_count += 1
                
                text_content.append("")  # Add space between sheets
            
            workbook.close()
            return '\n'.join(text_content) if text_content else None
            
        except Exception as e:
            print(f"Error extracting Excel: {str(e)}")
            return None
    
    def _extract_from_text(self, file: FileStorage) -> Optional[str]:
        """Extract text from plain text or CSV file"""
        try:
            # Try to decode as UTF-8 first, then fall back to other encodings
            content = file.read()
            
            encodings = ['utf-8', 'utf-8-sig', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    text = content.decode(encoding)
                    
                    # If it's a CSV file, try to parse it properly
                    if file.filename.lower().endswith('.csv'):
                        return self._parse_csv_content(text)
                    else:
                        return text
                        
                except UnicodeDecodeError:
                    continue
            
            # If all encodings fail, return None
            return None
            
        except Exception as e:
            print(f"Error extracting text: {str(e)}")
            return None
    
    def _parse_csv_content(self, csv_text: str) -> str:
        """Parse CSV content and convert to readable text"""
        try:
            text_content = []
            csv_reader = csv.reader(io.StringIO(csv_text))
            
            row_count = 0
            for row in csv_reader:
                if row_count >= self.max_rows:
                    text_content.append("... (truncated due to size limit)")
                    break
                
                # Filter out empty cells and join with separator
                row_values = [cell.strip() for cell in row if cell.strip()]
                if row_values:
                    text_content.append(' | '.join(row_values))
                
                row_count += 1
            
            return '\n'.join(text_content) if text_content else None
            
        except Exception as e:
            print(f"Error parsing CSV: {str(e)}")
            return csv_text  # Return original text if parsing fails
    
    def get_file_info(self, file: FileStorage) -> dict:
        """Get file information"""
        return {
            'filename': file.filename,
            'content_type': file.content_type,
            'size': len(file.read()) if hasattr(file, 'read') else None
        }